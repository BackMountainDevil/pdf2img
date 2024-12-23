import fitz  # PyMuPDF
from docx import Document
from docx.shared import Cm
from os import makedirs, listdir
from os.path import exists, splitext, join
from shutil import rmtree


def pdf_to_png(pdf_path, img_output_folder="tmp", dpi=300):
    """
    将PDF文件的每一页保存为PNG图片。文件名称为 页码.png

    :param pdf_path: PDF文件的路径
    :param img_output_folder: 保存PNG图片的文件夹路径
    :param dpi: 图片的分辨率（DPI），默认为300
    """
    # 如果输出文件夹不存在，创建它
    if not exists(img_output_folder):
        makedirs(img_output_folder)

    pdf_document = fitz.open(pdf_path)
    for page_number in range(len(pdf_document)):
        page = pdf_document[page_number]
        pix = page.get_pixmap(matrix=fitz.Matrix(dpi / 72, dpi / 72))  # 设置DPI
        output_path = f"{img_output_folder}/{page_number + 1}.png"
        pix.save(output_path)  # 保存为PNG图片
    print(f"所有页面已保存为PNG图片。一共{len(pdf_document)}页")
    pdf_document.close()  # 关闭PDF文件


def insert_images_to_word(image_folder, output_word_path):
    """
    将指定文件夹中的图片按顺序插入到Word文档中。

    :param image_folder: 包含PNG图片的文件夹路径
    :param output_word_path: 输出的Word文件路径
    """

    doc = Document()
    sections = doc.sections  # 设置页面为A4大小，常规页边距
    width = 21
    height = 29.7
    left_right_marin = 3.18
    top_bottom_margin = 2.54
    for section in sections:
        section.page_width = Cm(width)  # A4宽度：21厘米
        section.page_height = Cm(height)  # A4高度：29.7厘米
        section.left_margin = Cm(left_right_marin)  # 左边距
        section.right_margin = Cm(left_right_marin)  # 右边距
        section.top_margin = Cm(top_bottom_margin)  # 上边距
        section.bottom_margin = Cm(top_bottom_margin)  # 下边距

    # 获取文件夹中的所有图片文件
    image_files = [f for f in listdir(image_folder) if f.endswith(".png")]
    image_files.sort(key=lambda x: int(x.split(".")[0]))  # 按页码大小排序，确保顺序正确

    # 遍历图片文件并插入到Word文档中
    img_width = width - 2 * left_right_marin
    img_height = height - 2 * top_bottom_margin
    for image_file in image_files:
        image_path = join(image_folder, image_file)
        # 插入图片并缩放到页面大小
        doc.add_picture(image_path, width=Cm(img_width), height=Cm(img_height))
        doc.paragraphs[-1].alignment = 1  # 居中对齐
        # doc.add_paragraph()  # 添加一个空段落，避免图片重叠
    doc.save(output_word_path)  # 保存Word文档
    print(f"图片已成功插入到Word文档：{output_word_path}")


def pdf_to_word(pdf_path, output_word_path, img_output_folder="tmp", dpi=300):
    """
    将PDF文件的每一页保存为PNG图片，同时按顺序插入到Word文档中。

    :param pdf_path: PDF文件的路径
    :param img_output_folder: 保存PNG图片的文件夹路径
    :param output_word_path: 输出的Word文件路径

    结果上和下面两行差不多，效率高一点点，但是哈希校验结果不相同，但看不出差别在哪
    pdf_to_png(pdf_path, img_output_folder)
    insert_images_to_word(img_output_folder, output_word_path)
    """
    doc = Document()
    sections = doc.sections
    width = 21
    height = 29.7
    left_right_marin = 3.18
    top_bottom_margin = 2.54
    img_width = width - 2 * left_right_marin
    img_height = height - 2 * top_bottom_margin
    for section in sections:
        section.page_width = Cm(width)
        section.page_height = Cm(height)
        section.left_margin = Cm(left_right_marin)
        section.right_margin = Cm(left_right_marin)
        section.top_margin = Cm(top_bottom_margin)
        section.bottom_margin = Cm(top_bottom_margin)

    if not exists(img_output_folder):
        makedirs(img_output_folder)

    pdf_document = fitz.open(pdf_path)
    for page_number in range(len(pdf_document)):
        page = pdf_document[page_number]
        pix = page.get_pixmap(matrix=fitz.Matrix(dpi / 72, dpi / 72))
        output_path = f"{img_output_folder}/{page_number + 1}.png"
        pix.save(output_path)
        doc.add_picture(output_path, width=Cm(img_width), height=Cm(img_height))
        doc.paragraphs[-1].alignment = 1  # 居中对齐
    print(f"所有页面已保存为PNG图片。一共{len(pdf_document)}页")
    pdf_document.close()
    doc.save(output_word_path)
    print(f"图片已成功插入到Word文档：{output_word_path}")


def pdf_to_word_without_img(
    pdf_path, output_word_path: str = None, img_output_folder="tmp", dpi=300
):
    """
    将PDF页面的每一页插入到Word文档中。

    :param pdf_path: PDF文件的路径
    :param img_output_folder: 保存PNG图片的文件夹路径
    :param output_word_path: 输出的Word文件路径
    """
    if output_word_path is None:
        output_word_path = splitext(pdf_path)[0] + ".docx"  # 输出的Word文件路径
    else:
        if exists(output_word_path):
            assert (
                False
            ), f"输出文件 {output_word_path} 已存在，请删除或为输出文件取个别名"
    if exists(img_output_folder):
        assert (
            False
        ), f"图片文件夹 {img_output_folder} 已存在，请删除或为临时文件夹取个别名"
    else:
        pdf_to_word(pdf_path, output_word_path, img_output_folder, dpi)
        rmtree(img_output_folder)


def pdfs_to_words_without_img(
    pdf_dir, output_word_path: str = None, img_output_folder="tmp", dpi=300
):
    """
    将文件夹中的所有PDF文件，每个pdf转成一个Word文档。

    :param pdf_dir: PDF文件夹路径
    :param img_output_folder: 保存PNG图片的临时文件夹路径
    :param output_word_path: 输出的Word文件路径
    """

    files = listdir(pdf_dir)
    if not files:
        print(f"文件夹 {pdf_dir} 为空，请检查路径是否正确")
        return
    for filename in files:
        if filename.endswith(".pdf"):
            pdf_path = join(pdf_dir, filename)
            pdf_to_word_without_img(pdf_path, output_word_path, img_output_folder, dpi)


if __name__ == "__main__":
    pdf_path = "论文.pdf"  # 替换为你的PDF文件路径
    img_output_folder = "tmp"  # 替换为你想保存图片的文件夹路径
    output_word_path = "out.docx"  # 替换为你想保存的Word文件路径

    # pdf_to_png(pdf_path, img_output_folder)
    # insert_images_to_word(img_output_folder, output_word_path)
    # pdf_to_word(pdf_path, output_word_path)   # 等效于上面两行

    # pdf_to_word_without_img(pdf_path, )
    pdfs_to_words_without_img(
        "论文合集",
    )
